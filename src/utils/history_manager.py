# src/utils/history_manager.py
from __future__ import annotations
import io
from typing import List, Dict, Tuple
from docx import Document

def build_history_text(history: List[Dict]) -> str:
    """Renderiza o histórico completo em texto plano, preservando a ordem."""
    lines = []
    for m in history:
        role = m.get("role", "").capitalize()
        content = m.get("content", "")
        if not content:
            continue
        lines.append(f"{role}:\n{content}\n")
    return "\n".join(lines).strip()

def chunk_history_dynamic(
    history: List[Dict],
    max_chars_per_chunk: int = 8000,
    overlap: int = 800,
) -> Tuple[str, int, int, list[str]]:
    """
    Divide o histórico em blocos (chunking) para preservar sentido e evitar
    estouro de contexto. NÃO resume nem corta frases deliberadamente.
    Retorna:
      - texto_unido: concatenação dos chunks (com cabeçalho por parte)
      - total_chars: tamanho total do histórico
      - n_chunks: número de partes
      - chunks: lista com o conteúdo de cada parte
    """
    text = build_history_text(history)
    if not text:
        return "", 0, 0, []

    n = len(text)
    chunks, start = [], 0
    while start < n:
        end = min(start + max_chars_per_chunk, n)
        chunk = text[start:end]
        chunks.append(chunk)
        if end >= n:
            break
        start = max(0, end - overlap)

    parts = []
    for i, c in enumerate(chunks, 1):
        parts.append(f"### Histórico (parte {i})\n{c}")
    return "\n\n".join(parts), n, len(chunks), chunks

def export_history_to_txt(history: List[Dict]) -> io.BytesIO:
    """Exporta o histórico completo para um arquivo TXT em memória."""
    data = build_history_text(history).encode("utf-8")
    buf = io.BytesIO(data)
    buf.name = "chat_history.txt"
    return buf

def export_history_to_docx(history: List[Dict]) -> io.BytesIO:
    """Exporta o histórico completo para DOCX em memória."""
    doc = Document()
    doc.add_heading("Histórico de Conversa", level=1)
    for m in history:
        role = m.get("role", "").capitalize()
        content = m.get("content", "")
        if not content:
            continue
        doc.add_paragraph(role + ":", style="Heading 2")
        doc.add_paragraph(content)
        doc.add_paragraph("")  # espaçamento
    output = io.BytesIO()
    doc.save(output)
    output.seek(0)
    output.name = "chat_history.docx"
    return output
